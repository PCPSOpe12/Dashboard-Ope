#!/usr/bin/env python3
"""
Génère les documents (grille de risques .xlsx + déclaration .docx)
à partir des requêtes déposées dans data/requests/
"""
import json, shutil, os, re
from pathlib import Path
from datetime import datetime

try:
    import openpyxl
    from docx import Document
except ImportError as e:
    print(f"Dépendance manquante: {e}")
    import sys; sys.exit(1)

P2_FR = {'faible':'faible','modere':'modéré','moyen':'moyen','eleve':'élevé'}

def fill_grille(data, dst):
    src = Path('assets/grille_risques_template.xlsx')
    if not src.exists():
        print(f"  Template introuvable: {src}"); return False
    shutil.copy(src, dst)
    wb = openpyxl.load_workbook(dst)
    ws = wb.active
    ws['C5'] = data.get('nom','')
    ws['C6'] = data.get('date_lieu','')
    ws['C8'] = data.get('p1', 0)
    ws['C9'] = P2_FR.get(data.get('p2','faible'), 'faible')
    ws['C10'] = P2_FR.get(data.get('e1','faible'), 'faible')
    ws['C11'] = P2_FR.get(data.get('e2','faible'), 'faible')
    wb.save(dst)
    return True

def replace_in_para(para, old, new):
    full = para.text
    if old not in full: return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]: run.text = ''
    return True

def fill_declaration(data, dst):
    src = Path('assets/declaration_template.docx')
    if not src.exists():
        print(f"  Template introuvable: {src}"); return False
    shutil.copy(src, dst)
    doc = Document(dst)

    nom       = data.get('nom','')
    org_nom   = data.get('org_nom','')
    contact   = data.get('contact','')
    qualite   = data.get('qualite','')
    adresse   = data.get('adresse','')
    tel       = data.get('tel','')
    email     = data.get('email','')
    lieu      = data.get('lieu','')
    commune   = data.get('commune','')
    dept      = data.get('dept','')
    date      = data.get('date','')
    h_deb     = data.get('heure_debut','')
    h_fin     = data.get('heure_fin','')
    p1        = str(data.get('p1',0))
    i_val     = str(data.get('i','0')).replace('.',',')
    ris_val   = str(data.get('ris','0')).replace('.',',')
    nb_sr     = str(data.get('nbSR',0))
    nb_cp     = str(data.get('nbCP',1))
    nb_pse2   = str(data.get('nbPSE2',0))
    nb_pse1   = str(data.get('nbPSE1',0))
    nb_veh    = str(data.get('nbVeh',1))

    replacements = {
        'Francis Labardin': contact,
        'Responsable administratif': qualite,
        '5  allée du renard, 77420 Champs-sur-Marne': adresse,
        '06 62 41 34 07': tel,
        'francis.labardin@outlook.com': email,
        'Passage de grade ': nom + ' ',
        'Passage de grade': nom,
        'Dojo David Douillet - 1 allée Georges Clémenceau': lieu,
        'Savigny-sur-Orge': commune,
        '91600': dept,
        '08/02/2026': date,
        '8h30': h_deb,
        '11h30': h_fin,
        ': 15': ': ' + p1,
        ': 0,75': ': ' + i_val,
        ': 0,01125': ': ' + ris_val,
        'Intervenants-secouristes\xa0: 2': f'Intervenants-secouristes\xa0: {nb_sr}',
        'Chef de Poste : 1': f'Chef de Poste : {nb_cp}',
        'PSE2 : 0': f'PSE2 : {nb_pse2}',
        'PSE1 : 1': f'PSE1 : {nb_pse1}',
    }

    # Organisateur
    for para in doc.paragraphs:
        if para.text.strip().startswith('Organisateur'):
            for run in para.runs:
                if 'Organisateur' in run.text and not run.text.strip().endswith(org_nom):
                    run.text = 'Organisateur : ' + org_nom
                    break

    for para in doc.paragraphs:
        for old, new in replacements.items():
            replace_in_para(para, old, new)

    # Véhicules
    for para in doc.paragraphs:
        if 'Nombre de véhicules' in para.text:
            replace_in_para(para, '\t1 ', f'\t{nb_veh} ')
            replace_in_para(para, '\t1\t', f'\t{nb_veh}\t')

    doc.save(dst)
    return True

def process_requests():
    req_dir = Path('data/requests')
    out_dir = Path('data/generated')
    out_dir.mkdir(parents=True, exist_ok=True)

    if not req_dir.exists():
        print("Pas de répertoire data/requests — rien à faire")
        return

    for req_file in sorted(req_dir.glob('*.json')):
        print(f"Traitement: {req_file.name}")
        try:
            data = json.loads(req_file.read_text(encoding='utf-8'))
            stem = req_file.stem  # ex: grille_1001_2026

            if stem.startswith('grille_'):
                dst = out_dir / (stem + '.xlsx')
                ok = fill_grille(data, dst)
                print(f"  {'✓' if ok else '✗'} Grille → {dst}")

            elif stem.startswith('decl_'):
                dst = out_dir / (stem + '.docx')
                ok = fill_declaration(data, dst)
                print(f"  {'✓' if ok else '✗'} Déclaration → {dst}")

        except Exception as e:
            print(f"  Erreur {req_file.name}: {e}")

if __name__ == '__main__':
    process_requests()
    print("Génération terminée.")
